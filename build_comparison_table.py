import json
from docx import Document


def build_combined_data(all_llms_data):

    combined_dict = {}
    all_llm_names = set()

    for llm_obj in all_llms_data:
        llm_name = llm_obj["name"]
        all_llm_names.add(llm_name)

        llm_extractions = llm_obj["data"]["LLM_extractions"]
        for extraction in llm_extractions:
            idx = extraction["Index"]
            paragraph_text = extraction["text"]
            bel_results = extraction["Results"]  # array of {bel_statement, evidence}

            if idx not in combined_dict:
                combined_dict[idx] = {
                    "text": paragraph_text,
                    "items": []
                }

            for r in bel_results:
                combined_dict[idx]["items"].append({
                    "llm": llm_name,
                    "statement": r["bel_statement"],
                    "evidence": r["evidence"]
                })

    return {
        "combined": combined_dict,
        "all_llm_names": all_llm_names
    }


def build_narrow_tables_by_evidence(combined_info, docx_filename="comparison_by_evidence.docx"):
    """
    Creates a 2-col table for each paragraph:
      - First row (merged) for "Paragraph X: <text>"
      - Then one row per unique evidence.  (left=Evidence, right=bullet statements)
      - Also ensures we add a row for any LLM with no statements for that paragraph.
    """

    doc = Document()

    doc.add_heading("LLM BEL Comparison (Grouped by Evidence)", level=1)
    doc.add_paragraph(
        "Each paragraph has a 2-column table. The first row shows the paragraph text. "
        "Then we list each unique evidence once, with bullet points for each LLM's statements. "
        "If an LLM has no statements for this paragraph, we add a row that says 'No statements.'"
    )

    # Extract the actual data
    combined_dict = combined_info["combined"]
    all_llm_names = combined_info["all_llm_names"]

    # Sort paragraphs by numeric index
    for idx in sorted(combined_dict.keys(), key=lambda x: int(x)):
        paragraph_text = combined_dict[idx]["text"]
        items = combined_dict[idx]["items"]  # list of dict: {llm, statement, evidence}

        # Build a set of LLMs that actually appear in this paragraph
        llms_in_this_paragraph = set(i["llm"] for i in items)

        # Create a new table for this paragraph
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Shading Accent 1'

        # Row 1: Merged cell with paragraph text
        hdr_row = table.rows[0]
        hdr_cell = hdr_row.cells[0]
        hdr_cell.merge(hdr_row.cells[1])
        hdr_cell.text = f"Paragraph {idx}:\n{paragraph_text}"

        # Group items by evidence
        evidence_map = {}
        for it in items:
            ev_text = it["evidence"].strip()
            if not ev_text:
                ev_text = "[No Evidence]"
            if ev_text not in evidence_map:
                evidence_map[ev_text] = []
            evidence_map[ev_text].append(it)

        # Add one row for each unique evidence
        for ev_text, ev_items in evidence_map.items():
            row_cells = table.add_row().cells
            # Left col: Evidence text
            row_cells[0].text = f"Evidence:\n{ev_text}"

            # Right col: bullet lines <LLM in bold>: <statement in italics>
            for eitem in ev_items:
                p = row_cells[1].add_paragraph(style="List Bullet")
                run_llm = p.add_run(f"{eitem['llm']}: ")
                run_llm.bold = True
                run_stmt = p.add_run(eitem["statement"])
                run_stmt.italic = True

        #
        # For each LLM that did NOT appear in this paragraph's items, we add a row:
        # "No statements."
        #
        missing_llms = all_llm_names - llms_in_this_paragraph
        if missing_llms:
            for m_llm in missing_llms:
                row_cells = table.add_row().cells
                row_cells[0].text = "No Evidence"  # or you could say "No statements"
                # Right column: bullet with "m_llm: No statements"
                p = row_cells[1].add_paragraph(style="List Bullet")
                run_llm = p.add_run(f"{m_llm}: ")
                run_llm.bold = True
                run_stmt = p.add_run("No statements.")
                run_stmt.italic = True

        doc.add_paragraph()  # blank line after table

    doc.save(docx_filename)
    print(f"Saved docx to {docx_filename}")


if __name__ == "__main__":
    #
    # 1) Load multiple LLM JSON outputs
    #
    with open("min_prompt_openai.json", "r") as f:
        llm1_data = json.load(f)

    with open("min_prompt_claude3.5.json", "r") as f:
        llm2_data = json.load(f)

    with open("min_prompt_claude37.json", "r") as f:
        llm3_data = json.load(f)

    all_llms_data = [
        {"name": "gpt-4o",     "data": llm1_data},
        {"name": "claude3.5", "data": llm2_data},
        {"name": "claude3.7", "data": llm3_data},
    ]

    #
    # 2) Merge them
    #
    combined_info = build_combined_data(all_llms_data)

    #
    # 3) Build docx with 2-col table, grouping by evidence, bold LLM, italic statements,
    #    and also a row for any LLM that has no statements in that paragraph.
    #
    build_narrow_tables_by_evidence(combined_info, "comparison_by_evidence.docx")
